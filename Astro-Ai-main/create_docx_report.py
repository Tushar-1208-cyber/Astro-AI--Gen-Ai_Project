
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def create_report():
    print("Generating Final Report Word Document...")
    doc = Document()
    
    # Title
    title = doc.add_heading('Astro-Ai: A Hybrid Intelligent Educational Suite', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # --- DECLARATION ---
    doc.add_heading('Declaration by the Candidates', level=1)
    doc.add_paragraph('We hereby declare that the project entitled "Astro-Ai: A Hybrid Intelligent Educational Suite" has been carried out to fulfil the partial requirements for completion of the core-elective course Generative AI and LLMs offered in the 6th Semester of the Bachelor of Technology (B.Tech) program in the Department of Computer Science and Engineering during AY-2025-26 (even semester). This experimental work has been carried out by us and submitted to the course instructor Dr. Soharab Hossain Shaikh. Due acknowledgments have been made in the text of the project to all other materials used. This project has been prepared in full compliance with the requirements and constraints of the prescribed curriculum.')
    doc.add_page_break()

    # --- ABSTRACT ---
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph('The rapid advancement of Artificial Intelligence has transformed the educational landscape, yet challenges remain regarding high inference costs, data privacy, latency, and language barriers in rural and resource-constrained environments. This project presents Astro-Ai, an Intelligent Educational Suite designed to democratize personalized learning through locally deployed AI.\n\nUnlike traditional cloud-dependent systems, Astro-Ai operates entirely on an edge-computing architecture utilizing a custom fine-tuned Large Language Model (TinyLlama-1.1B). To achieve domain specialization and robust multilingual support (English, Hindi, Bengali, Bhojpuri), the model was extensively trained using Parameter-Efficient Fine-Tuning (QLoRA) on a custom-curated educational dataset. Furthermore, to mitigate the critical issue of AI hallucination, the system integrates Retrieval-Augmented Generation (RAG) utilizing ChromaDB as a vector store, ensuring contextually grounded and factually accurate responses without relying on external internet APIs.\n\nThe full-stack application, built on Next.js and FastAPI, features specialized modules including an interactive AI Tutor, automated Lesson Maker, and AI-powered Smart Attendance. Quantitative evaluations demonstrate a significant performance leap over generic pre-trained models, achieving an overall educational accuracy of 85% with optimized inference speeds via NVIDIA CUDA acceleration. The proposed framework successfully establishes a highly scalable, privacy-first, and cost-efficient educational ecosystem that runs entirely independent of third-party API keys.')
    doc.add_page_break()

    # --- TECH STACK ---
    doc.add_heading('Technology Stack', level=1)
    doc.add_paragraph('• Frontend: Next.js 15, React, Tailwind CSS\n• Backend API: FastAPI (Python), Node.js\n• AI Frameworks: PyTorch, Hugging Face Transformers, PEFT (QLoRA)\n• Local Model: TinyLlama-1.1B\n• Databases: MongoDB (User Data), ChromaDB (Vector Store)\n• Hardware Acceleration: NVIDIA CUDA (RTX 4050)')
    doc.add_page_break()

    # --- INTRODUCTION ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph('The integration of Artificial Intelligence (AI) in education has shown immense potential to revolutionize how students learn and teachers instruct. However, the current paradigm heavily relies on massive, cloud-based LLMs. While powerful, these generic models introduce significant challenges when deployed in real-world educational settings due to internet dependency, API costs, and data privacy risks.\n\nAstro-Ai addresses these challenges by acting as an entirely localized, privacy-first educational suite. By leveraging Edge AI and RAG, it delivers highly accurate, multilingual K-12 education entirely offline on consumer hardware.')
    
    # --- PROBLEM STATEMENT ---
    doc.add_heading('2. Problem Statement', level=1)
    doc.add_paragraph('The primary problems addressed by this project are:\n1. High Infrastructure Costs: Cloud APIs require constant internet and recurring fees.\n2. Data Privacy: Transmitting student data to third-party servers poses security risks.\n3. Language Barrier: Generic LLMs struggle with regional languages like Hindi and Bhojpuri.\n4. AI Hallucination: Generic models often generate factually incorrect educational info.')
    
    # --- LITERATURE REVIEW ---
    doc.add_heading('3. Literature Review', level=1)
    doc.add_paragraph('Recent research has pivoted towards Small Language Models (SLMs) and Edge AI to solve latency and privacy issues. The introduction of QLoRA by Dettmers et al. (2023) allowed massive models to be fine-tuned on consumer-grade GPUs. Furthermore, Lewis et al. (2020) proposed Retrieval-Augmented Generation (RAG) to solve AI hallucinations by combining parametric memory with vector databases (like ChromaDB). Astro-Ai combines these extensively researched methodologies into a single, cohesive educational application.')
    doc.add_page_break()

    # --- METHODOLOGY ---
    doc.add_heading('4. Methodology', level=1)
    doc.add_paragraph('The development of Astro-Ai followed a systematic 5-phase pipeline, shifting from raw data collection to full-stack deployment.')
    
    if os.path.exists('methodology_flow_diag.png'):
        doc.add_picture('methodology_flow_diag.png', width=Inches(6.0))
        p = doc.add_paragraph('Fig 1: Astro-Ai Project Methodology & Development Flow')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading('4.1 System Architecture', level=2)
    doc.add_paragraph('The system utilizes a Next.js frontend and FastAPI backend. To eliminate cloud dependency, the core AI engine runs locally on an NVIDIA CUDA GPU, utilizing ChromaDB for RAG context retrieval before generating responses.')
    
    if os.path.exists('local_architecture_diag.png'):
        doc.add_picture('local_architecture_diag.png', width=Inches(6.0))
        p = doc.add_paragraph('Fig 2: Astro-Ai Local-First System Architecture')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading('4.2 Dataset & Splitting', level=2)
    doc.add_paragraph('A custom multilingual educational dataset was curated and split using an 80:10:10 ratio to ensure unbiased training and evaluation.')
    
    if os.path.exists('dataset_split_chart.png'):
        doc.add_picture('dataset_split_chart.png', width=Inches(5.0))
        p = doc.add_paragraph('Fig 3: Dataset Split Strategy (80/10/10)')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # --- RESULTS & EVALUATION ---
    doc.add_heading('5. Results & Evaluation', level=1)

    doc.add_heading('5.1 Hardware Acceleration (CPU vs GPU)', level=2)
    doc.add_paragraph('Deploying the model required substantial computational power. Utilizing NVIDIA CUDA cores provided a massive performance leap over traditional CPU inference.')
    if os.path.exists('gpu_vs_cpu_comparison.png'):
        doc.add_picture('gpu_vs_cpu_comparison.png', width=Inches(5.5))
        p = doc.add_paragraph('Fig 4: Performance Leap: CPU vs CUDA GPU')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading('5.2 Baseline Comparison (Base vs Fine-Tuned)', level=2)
    doc.add_paragraph('The QLoRA fine-tuning process yielded dramatic improvements, especially in JSON formatting and regional language comprehension.')
    if os.path.exists('base_vs_finetuned_comparison.png'):
        doc.add_picture('base_vs_finetuned_comparison.png', width=Inches(5.5))
        p = doc.add_paragraph('Fig 5: Base Model vs Fine-Tuned Astro-Ai')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading('5.3 Cloud Prompting vs Local Fine-Tuning', level=2)
    if os.path.exists('prompt_vs_finetune_comparison.png'):
        doc.add_picture('prompt_vs_finetune_comparison.png', width=Inches(5.5))
        p = doc.add_paragraph('Fig 6: Cloud GPT-4 Prompting vs Local SFT')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_heading('5.4 Quantitative Metrics', level=2)
    if os.path.exists('evaluation_metrics_summary.png'):
        doc.add_picture('evaluation_metrics_summary.png', width=Inches(5.5))
        p = doc.add_paragraph('Fig 7: Final Quantitative Evaluation Metrics')
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()

    # --- UI FEATURES ---
    doc.add_heading('6. Key Features and User Interface', level=1)
    doc.add_paragraph('The Astro-Ai frontend features role-based dashboards.')
    
    doc.add_heading('6.1 Interactive AI Tutor', level=2)
    p = doc.add_paragraph('[INSERT SCREENSHOT HERE: ai_tutor_chat_ui.png]')
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = docx.shared.RGBColor(255, 0, 0) if hasattr(docx, 'shared') else None

    doc.add_heading('6.2 Automated Lesson Maker', level=2)
    p = doc.add_paragraph('[INSERT SCREENSHOT HERE: teacher_lesson_maker_ui.png]')
    p.runs[0].bold = True

    doc.add_heading('6.3 Multilingual Support', level=2)
    p = doc.add_paragraph('[INSERT SCREENSHOT HERE: multilingual_toggle_ui.png]')
    p.runs[0].bold = True

    doc.add_heading('6.4 Smart Attendance System', level=2)
    p = doc.add_paragraph('[INSERT SCREENSHOT HERE: smart_attendance_dashboard.png]')
    p.runs[0].bold = True

    doc.add_page_break()

    # --- CONCLUSION ---
    doc.add_heading('7. Conclusion & Future Scope', level=1)
    doc.add_paragraph('Astro-Ai successfully demonstrates the viability of Edge AI in democratizing education. By combining a QLoRA fine-tuned TinyLlama-1.1B model with ChromaDB-powered RAG, the project successfully eliminates cloud-API costs, protects student privacy, and neutralizes AI hallucination. Future scope includes scaling the application to mobile NPUs and integrating Voice-to-Voice regional interaction.')

    output_path = r'C:\Users\Laptop\OneDrive\Desktop\Astro_Ai_Final_Report_V2.docx'
    doc.save(output_path)
    print(f"✅ Success! Report generated at: {output_path}")

if __name__ == "__main__":
    try:
        import docx
    except ImportError:
        print("Installing python-docx...")
        os.system("pip install python-docx")
        import docx
    
    create_report()
